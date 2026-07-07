"""
Docs GPT engine — document extraction, corpus selection, and grounded Q&A
via the Claude Citations API.

Answers come ONLY from uploaded documents. Every supported claim carries
verbatim cited spans (char_location) that index 1:1 into the stored
extracted text, so the UI can highlight the exact passage.

Multi-user: each browser gets its own WORKSPACE (validated id -> its own
folder under data/ws/<id>/). Local single-user installs keep the legacy
flat layout under data/ as the "default" workspace, so existing libraries
keep working unchanged.

API keys resolve per request: pasted key (BYO) -> ANTHROPIC_API_KEY env
-> ~/.docsgpt/config.json. Keys are never logged and never stored by the
server in online mode.
"""

import io
import json
import os
import re
import shutil
import sys
import time
import uuid
from pathlib import Path

from dotenv import load_dotenv

import anthropic
from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

load_dotenv()
_EXTRA_ENV = os.getenv("DOCSGPT_ENV_FILE", "")
if _EXTRA_ENV and Path(_EXTRA_ENV).exists():
    load_dotenv(_EXTRA_ENV)

# Frozen (PyInstaller) builds unpack to a temp dir that is wiped on exit, so
# bundled assets are read from there but user data must live in the home dir.
FROZEN = bool(getattr(sys, "frozen", False))
APP_DIR = Path(getattr(sys, "_MEIPASS", "")) if FROZEN else Path(__file__).resolve().parent
DATA_DIR = (Path.home() / ".docsgpt" / "data") if FROZEN else APP_DIR / "data"
DEMO_DIR = APP_DIR / "demo-corpus"
CONFIG_PATH = Path.home() / ".docsgpt" / "config.json"

MODEL = os.getenv("DOCSGPT_MODEL", "claude-opus-4-8")
MAX_ANSWER_TOKENS = 16000
CORPUS_CHAR_BUDGET = 350_000
MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_DOCS_PER_WORKSPACE = 20
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
ALLOWED_STATUSES = {"current", "superseded"}

DEFAULT_WORKSPACE = "default"
WORKSPACE_RE = re.compile(r"^[a-f0-9][a-f0-9-]{7,63}$")

REFUSAL_SENTENCE = "Not found in your documents."

SYSTEM_PROMPT = f"""You are Docs GPT, a document-grounded assistant for SOPs and \
technical documents.
Rules you must never break:
1. Answer ONLY from the attached documents. Never use outside knowledge, even for facts
   you are certain of.
2. If the documents do not contain the answer, reply with exactly: "{REFUSAL_SENTENCE}"
   You may add ONE sentence naming the closest related content, clearly labeled as
   related but not an answer.
3. Keep answers short and practical. Every factual claim must come from the documents
   so it carries a citation.
4. Never guess values, limits, or procedure steps. A partial answer is fine if clearly
   scoped (e.g. "The documents specify X but not Y.").
5. The documents are the authority. Your answer is a pointer to them, not a replacement
   for reading the controlling procedure."""

DEMO_DOCS = [
    ("SOP-001_pump-startup_rev3.md",
     "SOP-001 Centrifugal Pump Startup & Shutdown", "3", "current"),
    ("SOP-014_confined-space-entry_rev5.md",
     "SOP-014 Confined Space Entry", "5", "current"),
    ("SOP-007_agitator-loto_rev2.md",
     "SOP-007 Agitator Lockout/Tagout", "2", "current"),
    ("SOP-007_agitator-loto_rev1_SUPERSEDED.md",
     "SOP-007 Agitator Lockout/Tagout (old)", "1", "superseded"),
]


class EngineError(Exception):
    """Raised for any engine failure the UI should show to the user."""


# ---------------- workspaces ----------------

def validate_workspace(workspace: str) -> str:
    """Return the workspace id if safe to use as a folder name, else raise."""
    if workspace == DEFAULT_WORKSPACE:
        return workspace
    if not WORKSPACE_RE.match(workspace or ""):
        raise EngineError("Invalid workspace id — refresh the page to get a new one.")
    return workspace


def _ws_root(workspace: str) -> Path:
    """Data root for a workspace (legacy flat layout for the default one)."""
    if workspace == DEFAULT_WORKSPACE:
        return DATA_DIR
    return DATA_DIR / "ws" / workspace


def _ws_paths(workspace: str) -> tuple[Path, Path, Path]:
    """Return (docs_dir, originals_dir, library_path) for a workspace."""
    root = _ws_root(validate_workspace(workspace))
    return root / "docs", root / "originals", root / "library.json"


def workspace_exists(workspace: str) -> bool:
    """True if the workspace already has a library file."""
    return _ws_paths(workspace)[2].exists()


def purge_stale_workspaces(max_idle_hours: float = 24.0) -> list[str]:
    """Delete non-default workspaces idle longer than the cutoff; return ids."""
    ws_container = DATA_DIR / "ws"
    if not ws_container.exists():
        return []
    cutoff = time.time() - max_idle_hours * 3600
    purged = []
    for folder in ws_container.iterdir():
        if not folder.is_dir():
            continue
        try:
            newest = max(
                (item.stat().st_mtime for item in folder.rglob("*") if item.is_file()),
                default=folder.stat().st_mtime,
            )
            if newest < cutoff:
                shutil.rmtree(folder)
                purged.append(folder.name)
        except OSError as exc:
            print(f"WARNING: could not purge workspace {folder.name} — {exc}")
    return purged


# ---------------- API key resolution (BYO) ----------------

def _read_config() -> dict:
    """Read ~/.docsgpt/config.json; missing or unreadable means empty."""
    if not CONFIG_PATH.exists():
        return {}
    try:
        with open(CONFIG_PATH, encoding="utf-8") as handle:
            return json.load(handle)
    except (OSError, json.JSONDecodeError):
        return {}


def env_key_present() -> bool:
    """True if ANTHROPIC_API_KEY is set in the server's environment."""
    return bool(os.getenv("ANTHROPIC_API_KEY", "").strip())


def saved_key_present() -> bool:
    """True if a key is saved in ~/.docsgpt/config.json."""
    return bool(str(_read_config().get("anthropic_api_key", "")).strip())


def save_api_key(key: str) -> None:
    """Persist the BYO key to ~/.docsgpt/config.json; empty key clears it."""
    try:
        CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
        config = _read_config()
        key = (key or "").strip()
        if key:
            config["anthropic_api_key"] = key
        else:
            config.pop("anthropic_api_key", None)
        with open(CONFIG_PATH, "w", encoding="utf-8") as handle:
            json.dump(config, handle)
    except OSError as exc:
        raise EngineError(f"Cannot write {CONFIG_PATH} — {exc}") from exc


def resolve_api_key(request_key: str | None = None) -> str:
    """Pasted key -> env -> saved config, in that order; raise if none."""
    key = (
        (request_key or "").strip()
        or os.getenv("ANTHROPIC_API_KEY", "").strip()
        or str(_read_config().get("anthropic_api_key", "")).strip()
    )
    if not key:
        raise EngineError(
            "No Anthropic API key — open Settings (⚙) and paste your key, "
            "or set ANTHROPIC_API_KEY."
        )
    return key


def _client(request_key: str | None = None) -> anthropic.Anthropic:
    """Build the Anthropic client with the resolved per-request key."""
    return anthropic.Anthropic(api_key=resolve_api_key(request_key))


# ---------------- library ----------------

def _ensure_dirs(workspace: str) -> None:
    """Create the workspace's data directories if they do not exist yet."""
    docs_dir, originals_dir, _ = _ws_paths(workspace)
    for directory in (docs_dir, originals_dir):
        directory.mkdir(parents=True, exist_ok=True)


def _read_library(workspace: str) -> list[dict]:
    """Load the document metadata list for a workspace."""
    library_path = _ws_paths(workspace)[2]
    if not library_path.exists():
        return []
    try:
        with open(library_path, encoding="utf-8") as handle:
            return json.load(handle)
    except (OSError, json.JSONDecodeError) as exc:
        raise EngineError(f"Cannot read library.json — {exc}") from exc


def _write_library(library: list[dict], workspace: str) -> None:
    """Persist the document metadata list for a workspace."""
    _ensure_dirs(workspace)
    library_path = _ws_paths(workspace)[2]
    try:
        with open(library_path, "w", encoding="utf-8") as handle:
            json.dump(library, handle, indent=2, ensure_ascii=False)
    except OSError as exc:
        raise EngineError(f"Cannot write library.json — {exc}") from exc


def list_documents(workspace: str = DEFAULT_WORKSPACE) -> list[dict]:
    """Return metadata for every document in a workspace."""
    return _read_library(workspace)


def _doc_text_path(doc_id: str, workspace: str) -> Path:
    """Path of the stored extracted text for a document."""
    return _ws_paths(workspace)[0] / f"{doc_id}.txt"


def get_document(doc_id: str, workspace: str = DEFAULT_WORKSPACE) -> dict:
    """Return one document's metadata plus its full extracted text."""
    for meta in _read_library(workspace):
        if meta["id"] == doc_id:
            text_path = _doc_text_path(doc_id, workspace)
            if not text_path.exists():
                raise EngineError(f"Stored text missing for document {doc_id}")
            try:
                with open(text_path, encoding="utf-8") as handle:
                    text = handle.read()
            except OSError as exc:
                raise EngineError(f"Cannot read document text — {exc}") from exc
            return {**meta, "text": text}
    raise EngineError(f"Document not found: {doc_id}")


def _extract_pdf(raw: bytes) -> str:
    """Extract text from a PDF, page by page, with page markers."""
    try:
        reader = PdfReader(io.BytesIO(raw))
    except PdfReadError as exc:
        raise EngineError(f"Cannot parse PDF — {exc}") from exc
    pages = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"[Page {number}]\n{text.strip()}")
    combined = "\n\n".join(pages).strip()
    if not combined.replace("[Page 1]", "").strip():
        raise EngineError(
            "No extractable text in this PDF — it may be a scanned image "
            "(OCR is not supported in v1)."
        )
    return combined


def _extract_docx(raw: bytes) -> str:
    """Extract paragraphs and simple table rows from a .docx file."""
    try:
        document = DocxDocument(io.BytesIO(raw))
    except (OSError, KeyError, ValueError) as exc:
        raise EngineError(f"Cannot parse .docx — {exc}") from exc
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(filename: str, raw: bytes) -> str:
    """Extract plain text from an uploaded file, dispatching on extension."""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise EngineError(
            f"Unsupported file type '{suffix}' — allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    if suffix == ".pdf":
        return _extract_pdf(raw)
    if suffix == ".docx":
        return _extract_docx(raw)
    text = raw.decode("utf-8", errors="replace").strip()
    if not text:
        raise EngineError("The uploaded file is empty.")
    return text


def add_document(  # pylint: disable=too-many-arguments
    filename: str,
    raw: bytes,
    *,
    title: str = "",
    rev: str = "",
    status: str = "current",
    workspace: str = DEFAULT_WORKSPACE,
) -> dict:
    """Extract, store, and register a new document; returns its metadata."""
    if len(raw) > MAX_UPLOAD_BYTES:
        raise EngineError("File exceeds the 15 MB upload limit.")
    if status not in ALLOWED_STATUSES:
        raise EngineError(f"Status must be one of: {', '.join(sorted(ALLOWED_STATUSES))}")
    library = _read_library(workspace)
    if workspace != DEFAULT_WORKSPACE and len(library) >= MAX_DOCS_PER_WORKSPACE:
        raise EngineError(
            f"Workspace is full ({MAX_DOCS_PER_WORKSPACE} documents max) — "
            "this is a demo space, not a document vault."
        )
    text = extract_text(filename, raw)
    _ensure_dirs(workspace)
    doc_id = uuid.uuid4().hex[:12]
    safe_name = Path(filename).name
    meta = {
        "id": doc_id,
        "title": title.strip() or Path(safe_name).stem,
        "filename": safe_name,
        "rev": rev.strip(),
        "status": status,
        "chars": len(text),
    }
    originals_dir = _ws_paths(workspace)[1]
    try:
        with open(_doc_text_path(doc_id, workspace), "w", encoding="utf-8") as handle:
            handle.write(text)
        with open(originals_dir / f"{doc_id}_{safe_name}", "wb") as handle:
            handle.write(raw)
    except OSError as exc:
        raise EngineError(f"Cannot store document — {exc}") from exc
    library.append(meta)
    _write_library(library, workspace)
    return meta


def set_status(doc_id: str, status: str, workspace: str = DEFAULT_WORKSPACE) -> dict:
    """Mark a document current or superseded; returns updated metadata."""
    if status not in ALLOWED_STATUSES:
        raise EngineError(f"Status must be one of: {', '.join(sorted(ALLOWED_STATUSES))}")
    library = _read_library(workspace)
    for meta in library:
        if meta["id"] == doc_id:
            meta["status"] = status
            _write_library(library, workspace)
            return meta
    raise EngineError(f"Document not found: {doc_id}")


def seed_demo_corpus(workspace: str = DEFAULT_WORKSPACE) -> int:
    """Load the synthetic demo SOPs into a workspace (idempotent by title)."""
    existing_titles = {doc["title"] for doc in list_documents(workspace)}
    added = 0
    for filename, title, rev, status in DEMO_DOCS:
        if title in existing_titles:
            continue
        path = DEMO_DIR / filename
        try:
            raw = path.read_bytes()
        except OSError as exc:
            raise EngineError(f"Cannot read demo file {filename} — {exc}") from exc
        add_document(filename, raw, title=title, rev=rev, status=status, workspace=workspace)
        added += 1
    return added


# ---------------- Q&A ----------------

def _question_terms(question: str) -> list[str]:
    """Lowercased keyword terms (3+ chars) from the question."""
    return [term for term in re.findall(r"[a-z0-9]+", question.lower()) if len(term) >= 3]


def score_document(question: str, text: str) -> int:
    """Simple keyword-overlap relevance score of a document for a question."""
    lowered = text.lower()
    return sum(lowered.count(term) for term in _question_terms(question))


def select_corpus(question: str, docs: list[dict]) -> tuple[list[dict], list[dict]]:
    """Pick documents within the corpus budget; returns (selected, dropped).

    If everything fits the char budget all docs go in. Otherwise docs are
    ranked by keyword overlap with the question and taken until the budget
    is spent — the dropped list is surfaced to the UI so the truncation is
    never silent.
    """
    total = sum(doc["chars"] for doc in docs)
    if total <= CORPUS_CHAR_BUDGET:
        return docs, []
    ranked = sorted(docs, key=lambda d: score_document(question, d["text"]), reverse=True)
    selected: list[dict] = []
    used = 0
    for doc in ranked:
        if used + doc["chars"] <= CORPUS_CHAR_BUDGET:
            selected.append(doc)
            used += doc["chars"]
    dropped = [doc for doc in ranked if doc not in selected]
    return selected, dropped


def number_citations(segments: list[dict]) -> list[dict]:
    """Assign stable [n] numbers across segments; dedupes identical spans.

    Mutates each segment's citations to add a "n" field and returns the
    flat, ordered list of unique citations for the evidence panel.
    """
    unique: list[dict] = []
    seen: dict[tuple, int] = {}
    for segment in segments:
        for cite in segment["citations"]:
            key = (cite["doc_id"], cite["start"], cite["end"])
            if key not in seen:
                seen[key] = len(unique) + 1
                unique.append(cite)
            cite["n"] = seen[key]
    return unique


def _build_content(question: str, selected: list[dict]) -> list[dict]:
    """Assemble the user-message content blocks: documents first, question last."""
    content: list[dict] = []
    for position, doc in enumerate(selected):
        rev_label = f"Rev {doc['rev']}" if doc["rev"] else "no rev"
        block = {
            "type": "document",
            "source": {"type": "text", "media_type": "text/plain", "data": doc["text"]},
            "title": f"{doc['title']} — {rev_label} ({doc['status']})",
            "citations": {"enabled": True},
        }
        if position == len(selected) - 1:
            block["cache_control"] = {"type": "ephemeral"}
        content.append(block)
    content.append({"type": "text", "text": question})
    return content


def _parse_segments(response, selected: list[dict]) -> list[dict]:
    """Map the response's cited text blocks onto our document metadata."""
    segments: list[dict] = []
    for block in response.content:
        if block.type != "text":
            continue
        cites = []
        for cite in block.citations or []:
            if cite.type != "char_location":
                continue
            doc = selected[cite.document_index]
            cites.append(
                {
                    "doc_id": doc["id"],
                    "doc_title": doc["title"],
                    "rev": doc["rev"],
                    "status": doc["status"],
                    "cited_text": cite.cited_text,
                    "start": cite.start_char_index,
                    "end": cite.end_char_index,
                }
            )
        segments.append({"text": block.text, "citations": cites})
    return segments


def ask(
    question: str,
    include_superseded: bool = False,
    workspace: str = DEFAULT_WORKSPACE,
    api_key: str | None = None,
) -> dict:
    """Answer a question grounded ONLY in the workspace's documents.

    Returns segments (text + citations), the numbered evidence list, which
    documents were searched/dropped, and a not_found flag for the refusal path.
    """
    question = (question or "").strip()
    if not question:
        raise EngineError("Question is empty.")
    docs = [get_document(meta["id"], workspace) for meta in _read_library(workspace)]
    if not include_superseded:
        docs = [doc for doc in docs if doc["status"] == "current"]
    if not docs:
        raise EngineError(
            "No documents available to search — upload a document first "
            "(or enable 'include superseded')."
        )
    selected, dropped = select_corpus(question, docs)

    try:
        response = _client(api_key).messages.create(
            model=MODEL,
            max_tokens=MAX_ANSWER_TOKENS,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": _build_content(question, selected)}],
        )
    except anthropic.AuthenticationError as exc:
        raise EngineError(
            "Claude API rejected the API key — check the key in Settings (⚙)."
        ) from exc
    except anthropic.RateLimitError as exc:
        raise EngineError("Claude API rate limit hit — wait a minute and retry.") from exc
    except anthropic.APIConnectionError as exc:
        raise EngineError(
            f"Cannot reach the Claude API — check the internet connection. ({exc})"
        ) from exc
    except anthropic.APIStatusError as exc:
        raise EngineError(f"Claude API error {exc.status_code} — {exc.message}") from exc

    if response.stop_reason == "refusal":
        raise EngineError("The model declined to answer this question (safety refusal).")

    segments = _parse_segments(response, selected)
    evidence = number_citations(segments)
    full_text = "".join(segment["text"] for segment in segments)
    cited_ids = {cite["doc_id"] for cite in evidence}
    return {
        "segments": segments,
        "evidence": evidence,
        "not_found": not evidence or REFUSAL_SENTENCE.rstrip(".") in full_text,
        "searched": [
            {"id": d["id"], "title": d["title"], "rev": d["rev"], "cited": d["id"] in cited_ids}
            for d in selected
        ],
        "dropped": [{"id": d["id"], "title": d["title"]} for d in dropped],
        "model": response.model,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
            "cache_read_input_tokens": getattr(response.usage, "cache_read_input_tokens", 0) or 0,
        },
    }
